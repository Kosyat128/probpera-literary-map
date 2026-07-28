#!/usr/bin/env python3
import html
import os
import json
import re
import time
from collections import OrderedDict
from pathlib import Path
from urllib.parse import quote

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

API = "https://ru.wikisource.org/w/api.php"
UA = "ProbPera-public-domain-anthology/1.0 (personal reading compilation)"
STORIES = [
  {
    "author": "Саки",
    "title": "Музыка на холме",
    "original": "The Music on the Hill"
  },
  {
    "author": "Саки",
    "title": "Фон",
    "original": "The Background"
  },
  {
    "author": "Саки",
    "title": "Пасхальное яйцо",
    "original": "The Easter Egg"
  },
  {
    "author": "Саки",
    "title": "Волки Чернограца",
    "original": "The Wolves of Cernogratz"
  },
  {
    "author": "Саки",
    "title": "Паутина",
    "original": "The Cobweb"
  },
  {
    "author": "Саки",
    "title": "Византийский омлет",
    "original": "The Byzantine Omelette"
  },
  {
    "author": "Саки",
    "title": "Волчица",
    "original": "The She-Wolf"
  },
  {
    "author": "Артур Конан Дойл",
    "title": "Бразильский кот",
    "original": "The Brazilian Cat"
  },
  {
    "author": "Артур Конан Дойл",
    "title": "Кожаная воронка",
    "original": "The Leather Funnel"
  },
  {
    "author": "Артур Конан Дойл",
    "title": "Исчезнувший экстренный поезд",
    "original": "The Lost Special"
  },
  {
    "author": "Артур Конан Дойл",
    "title": "Коричневая рука",
    "original": "The Brown Hand"
  },
  {
    "author": "Артур Конан Дойл",
    "title": "Ужас высот",
    "original": "The Horror of the Heights"
  },
  {
    "author": "Артур Конан Дойл",
    "title": "Ужас ущелья Блю-Джон",
    "original": "The Terror of Blue John Gap"
  },
  {
    "author": "Артур Конан Дойл",
    "title": "Лакированная шкатулка",
    "original": "The Japanned Box"
  },
  {
    "author": "Редьярд Киплинг",
    "title": "Домашний хирург",
    "original": "The House Surgeon"
  },
  {
    "author": "Редьярд Киплинг",
    "title": "В конце пути",
    "original": "At the End of the Passage"
  },
  {
    "author": "Редьярд Киплинг",
    "title": "В одной лодке",
    "original": "In the Same Boat"
  },
  {
    "author": "Редьярд Киплинг",
    "title": "Дом желаний",
    "original": "The Wish House"
  },
  {
    "author": "Редьярд Киплинг",
    "title": "Пёс Херви",
    "original": "The Dog Hervey"
  },
  {
    "author": "Редьярд Киплинг",
    "title": "Из уст в уста",
    "original": "By Word of Mouth"
  },
  {
    "author": "Редьярд Киплинг",
    "title": "Послание Даны Да",
    "original": "The Sending of Dana Da"
  },
  {
    "author": "Натаниэль Готорн",
    "title": "Великий карбункул",
    "original": "The Great Carbuncle"
  },
  {
    "author": "Натаниэль Готорн",
    "title": "Честолюбивый гость",
    "original": "The Ambitious Guest"
  },
  {
    "author": "Натаниэль Готорн",
    "title": "Тройная судьба",
    "original": "The Threefold Destiny"
  },
  {
    "author": "Натаниэль Готорн",
    "title": "Мастер прекрасного",
    "original": "The Artist of the Beautiful"
  },
  {
    "author": "Натаниэль Готорн",
    "title": "Эгоизм, или Змея в груди",
    "original": "Egotism; or, The Bosom Serpent"
  },
  {
    "author": "Натаниэль Готорн",
    "title": "Рождественский пир",
    "original": "The Christmas Banquet"
  },
  {
    "author": "Натаниэль Готорн",
    "title": "Снежный образ",
    "original": "The Snow-Image: A Childish Miracle"
  },
  {
    "author": "Роберт Льюис Стивенсон",
    "title": "Кривая Дженет",
    "original": "Thrawn Janet"
  },
  {
    "author": "Роберт Льюис Стивенсон",
    "title": "Весёлые молодцы",
    "original": "The Merry Men"
  },
  {
    "author": "Роберт Льюис Стивенсон",
    "title": "Клад Франшара",
    "original": "The Treasure of Franchard"
  },
  {
    "author": "Роберт Льюис Стивенсон",
    "title": "Маркхейм",
    "original": "Markheim"
  },
  {
    "author": "Роберт Льюис Стивенсон",
    "title": "Олалла",
    "original": "Olalla"
  },
  {
    "author": "Роберт Льюис Стивенсон",
    "title": "Похититель трупов",
    "original": "The Body Snatcher"
  },
  {
    "author": "Роберт Льюис Стивенсон",
    "title": "Сатанинская бутылка",
    "original": "The Bottle Imp"
  },
  {
    "author": "Чарльз Диккенс",
    "title": "Мистер Миннс и его кузен",
    "original": "Mr. Minns and His Cousin"
  },
  {
    "author": "Чарльз Диккенс",
    "title": "Чёрная вуаль",
    "original": "The Black Veil"
  },
  {
    "author": "Чарльз Диккенс",
    "title": "Гоблины, похитившие могильщика",
    "original": "The Goblins Who Stole a Sexton"
  },
  {
    "author": "Чарльз Диккенс",
    "title": "Суд по делу об убийстве",
    "original": "The Trial for Murder"
  },
  {
    "author": "Чарльз Диккенс",
    "title": "Рассказ дядюшки коммивояжёра",
    "original": "The Story of the Bagman’s Uncle"
  },
  {
    "author": "Чарльз Диккенс",
    "title": "Доктор Мэриголд",
    "original": "Doctor Marigold"
  },
  {
    "author": "Чарльз Диккенс",
    "title": "Выход в общество",
    "original": "Going into Society"
  },
  {
    "author": "Э. Т. А. Гофман",
    "title": "Советник Креспель",
    "original": "Rat Krespel"
  },
  {
    "author": "Э. Т. А. Гофман",
    "title": "Фермата",
    "original": "Die Fermate"
  },
  {
    "author": "Э. Т. А. Гофман",
    "title": "Майорат",
    "original": "Das Majorat"
  },
  {
    "author": "Э. Т. А. Гофман",
    "title": "Пустой дом",
    "original": "Das öde Haus"
  },
  {
    "author": "Э. Т. А. Гофман",
    "title": "Счастье игрока",
    "original": "Spielerglück"
  },
  {
    "author": "Э. Т. А. Гофман",
    "title": "Выбор невесты",
    "original": "Die Brautwahl"
  },
  {
    "author": "Э. Т. А. Гофман",
    "title": "Дож и догаресса",
    "original": "Doge und Dogaresse"
  },
  {
    "author": "Оскар Уайльд",
    "title": "Образцовый миллионер",
    "original": "The Model Millionaire"
  },
  {
    "author": "Оскар Уайльд",
    "title": "Сфинкс без загадки",
    "original": "The Sphinx Without a Secret"
  },
  {
    "author": "Оскар Уайльд",
    "title": "Портрет господина У. Х.",
    "original": "The Portrait of Mr. W. H."
  },
  {
    "author": "Оскар Уайльд",
    "title": "Кентервильское привидение",
    "original": "The Canterville Ghost"
  },
  {
    "author": "Оскар Уайльд",
    "title": "Соловей и роза",
    "original": "The Nightingale and the Rose"
  },
  {
    "author": "Оскар Уайльд",
    "title": "Замечательная ракета",
    "original": "The Remarkable Rocket"
  },
  {
    "author": "Оскар Уайльд",
    "title": "Преданный друг",
    "original": "The Devoted Friend"
  },
  {
    "author": "О. Генри",
    "title": "Дороги, которые мы выбираем",
    "original": "The Roads We Take"
  },
  {
    "author": "О. Генри",
    "title": "Квадратура круга",
    "original": "Squaring the Circle"
  },
  {
    "author": "О. Генри",
    "title": "Фараон и хорал",
    "original": "The Cop and the Anthem"
  },
  {
    "author": "О. Генри",
    "title": "Поросячья этика",
    "original": "The Ethics of Pig"
  },
  {
    "author": "О. Генри",
    "title": "Русские соболя",
    "original": "The Ransom of Red Chief"
  },
  {
    "author": "О. Генри",
    "title": "Трест, который лопнул",
    "original": "The Octopus Marooned"
  },
  {
    "author": "О. Генри",
    "title": "Коловращение жизни",
    "original": "The Whirligig of Life"
  },
  {
    "author": "Ги де Мопассан",
    "title": "Покойница",
    "original": "La Morte"
  },
  {
    "author": "Ги де Мопассан",
    "title": "Рука",
    "original": "La Main"
  },
  {
    "author": "Ги де Мопассан",
    "title": "Страх",
    "original": "La Peur"
  },
  {
    "author": "Ги де Мопассан",
    "title": "Гораций",
    "original": "Horla"
  },
  {
    "author": "Ги де Мопассан",
    "title": "Кто знает?",
    "original": "Qui sait?"
  },
  {
    "author": "Ги де Мопассан",
    "title": "Ожерелье",
    "original": "La Parure"
  },
  {
    "author": "Ги де Мопассан",
    "title": "На воде",
    "original": "Sur l’eau"
  },
  {
    "author": "Анатоль Франс",
    "title": "Прокуратор Иудеи",
    "original": "Le Procurateur de Judée"
  },
  {
    "author": "Анатоль Франс",
    "title": "Жонглёр Богоматери",
    "original": "Le Jongleur de Notre-Dame"
  },
  {
    "author": "Анатоль Франс",
    "title": "Мессир Гвидо Кавальканти",
    "original": "Messire Guido Cavalcanti"
  },
  {
    "author": "Анатоль Франс",
    "title": "Лесли Вуд",
    "original": "Leslie Wood"
  },
  {
    "author": "Анатоль Франс",
    "title": "Амикус и Целестин",
    "original": "Amicus et Célestin"
  },
  {
    "author": "Анатоль Франс",
    "title": "Легенда о святых Оливерии и Либеретте",
    "original": "La Légende des saintes Oliverie et Liberette"
  },
  {
    "author": "Анатоль Франс",
    "title": "Записки сельского врача",
    "original": "Mémoires d’un médecin de campagne"
  },
  {
    "author": "Сомерсет Моэм",
    "title": "Творческий импульс",
    "original": "The Creative Impulse"
  },
  {
    "author": "Сомерсет Моэм",
    "title": "Друг в беде",
    "original": "A Friend in Need"
  },
  {
    "author": "Сомерсет Моэм",
    "title": "Человек со шрамом",
    "original": "The Man with the Scar"
  },
  {
    "author": "Сомерсет Моэм",
    "title": "Маляр",
    "original": "The Painter"
  },
  {
    "author": "Сомерсет Моэм",
    "title": "Конец полёта",
    "original": "The End of the Flight"
  },
  {
    "author": "Сомерсет Моэм",
    "title": "Дюжина",
    "original": "The Round Dozen"
  },
  {
    "author": "Сомерсет Моэм",
    "title": "Завтрак",
    "original": "The Luncheon"
  },
  {
    "author": "Франц Кафка",
    "title": "Приговор",
    "original": "Das Urteil"
  },
  {
    "author": "Франц Кафка",
    "title": "В исправительной колонии",
    "original": "In der Strafkolonie"
  },
  {
    "author": "Франц Кафка",
    "title": "Сельский врач",
    "original": "Ein Landarzt"
  },
  {
    "author": "Франц Кафка",
    "title": "Голодарь",
    "original": "Ein Hungerkünstler"
  },
  {
    "author": "Франц Кафка",
    "title": "Нора",
    "original": "Der Bau"
  },
  {
    "author": "Франц Кафка",
    "title": "Братоубийство",
    "original": "Ein Brudermord"
  },
  {
    "author": "Франц Кафка",
    "title": "Старый лист",
    "original": "Ein altes Blatt"
  },
  {
    "author": "Кэтрин Мэнсфилд",
    "title": "Кукольный дом",
    "original": "The Doll’s House"
  },
  {
    "author": "Кэтрин Мэнсфилд",
    "title": "Незнакомец",
    "original": "The Stranger"
  },
  {
    "author": "Кэтрин Мэнсфилд",
    "title": "Жизнь мамаши Паркер",
    "original": "Life of Ma Parker"
  },
  {
    "author": "Кэтрин Мэнсфилд",
    "title": "Маленькая гувернантка",
    "original": "The Little Governess"
  },
  {
    "author": "Кэтрин Мэнсфилд",
    "title": "Нескромное путешествие",
    "original": "An Indiscreet Journey"
  },
  {
    "author": "Кэтрин Мэнсфилд",
    "title": "Принять постриг",
    "original": "Taking the Veil"
  },
  {
    "author": "Кэтрин Мэнсфилд",
    "title": "Ветер дует",
    "original": "The Wind Blows"
  },
  {
    "author": "Иван Тургенев",
    "title": "Собака",
    "original": ""
  },
  {
    "author": "Иван Тургенев",
    "title": "История лейтенанта Ергунова",
    "original": ""
  },
  {
    "author": "Иван Тургенев",
    "title": "Странная история",
    "original": ""
  },
  {
    "author": "Иван Тургенев",
    "title": "Стук… стук… стук!",
    "original": ""
  },
  {
    "author": "Иван Тургенев",
    "title": "Сон",
    "original": ""
  },
  {
    "author": "Иван Тургенев",
    "title": "Песнь торжествующей любви",
    "original": ""
  },
  {
    "author": "Иван Тургенев",
    "title": "Рассказ отца Алексея",
    "original": ""
  },
  {
    "author": "Николай Лесков",
    "title": "Привидение в Инженерном замке",
    "original": ""
  },
  {
    "author": "Николай Лесков",
    "title": "Неразменный рубль",
    "original": ""
  },
  {
    "author": "Николай Лесков",
    "title": "Пугало",
    "original": ""
  },
  {
    "author": "Николай Лесков",
    "title": "Зверь",
    "original": ""
  },
  {
    "author": "Николай Лесков",
    "title": "Старый гений",
    "original": ""
  },
  {
    "author": "Николай Лесков",
    "title": "Человек на часах",
    "original": ""
  },
  {
    "author": "Николай Лесков",
    "title": "Штопальщик",
    "original": ""
  },
  {
    "author": "Максим Горький",
    "title": "Дед Архип и Лёнька",
    "original": ""
  },
  {
    "author": "Максим Горький",
    "title": "Мальва",
    "original": ""
  },
  {
    "author": "Максим Горький",
    "title": "Коновалов",
    "original": ""
  },
  {
    "author": "Максим Горький",
    "title": "Проходимец",
    "original": ""
  },
  {
    "author": "Максим Горький",
    "title": "Супруги Орловы",
    "original": ""
  },
  {
    "author": "Максим Горький",
    "title": "Бывшие люди",
    "original": ""
  },
  {
    "author": "Максим Горький",
    "title": "Ошибка",
    "original": ""
  },
  {
    "author": "Иван Бунин",
    "title": "Кукушка",
    "original": ""
  },
  {
    "author": "Иван Бунин",
    "title": "Ночной разговор",
    "original": ""
  },
  {
    "author": "Иван Бунин",
    "title": "Худая трава",
    "original": ""
  },
  {
    "author": "Иван Бунин",
    "title": "Иоанн Рыдалец",
    "original": ""
  },
  {
    "author": "Иван Бунин",
    "title": "Петлистые уши",
    "original": ""
  },
  {
    "author": "Иван Бунин",
    "title": "Молодость и старость",
    "original": ""
  },
  {
    "author": "Иван Бунин",
    "title": "Книга",
    "original": ""
  },
  {
    "author": "Александр Куприн",
    "title": "Собачье счастье",
    "original": ""
  },
  {
    "author": "Александр Куприн",
    "title": "Брегет",
    "original": ""
  },
  {
    "author": "Александр Куприн",
    "title": "На разъезде",
    "original": ""
  },
  {
    "author": "Александр Куприн",
    "title": "Пустые дачи",
    "original": ""
  },
  {
    "author": "Александр Куприн",
    "title": "Сапсан",
    "original": ""
  },
  {
    "author": "Александр Куприн",
    "title": "Ю-ю",
    "original": ""
  },
  {
    "author": "Александр Куприн",
    "title": "Обида",
    "original": ""
  }
]

AUTHOR_ALIASES = {
    "Саки": ["Саки", "Гектор Хью Манро", "Манро"],
    "Артур Конан Дойл": ["Артур Конан Дойл", "Конан Дойл", "Дойл"],
    "Редьярд Киплинг": ["Редьярд Киплинг", "Киплинг"],
    "Натаниэль Готорн": ["Натаниэль Готорн", "Готорн"],
    "Роберт Льюис Стивенсон": ["Роберт Льюис Стивенсон", "Стивенсон"],
    "Чарльз Диккенс": ["Чарльз Диккенс", "Диккенс"],
    "Э. Т. А. Гофман": ["Э. Т. А. Гофман", "Гофман"],
    "Оскар Уайльд": ["Оскар Уайльд", "Уайльд"],
    "О. Генри": ["О. Генри", "О Генри"],
    "Ги де Мопассан": ["Ги де Мопассан", "Мопассан"],
    "Анатоль Франс": ["Анатоль Франс", "Франс"],
    "Сомерсет Моэм": ["Сомерсет Моэм", "Моэм"],
    "Франц Кафка": ["Франц Кафка", "Кафка"],
    "Кэтрин Мэнсфилд": ["Кэтрин Мэнсфилд", "Мэнсфилд"],
    "Иван Тургенев": ["Иван Тургенев", "Тургенев"],
    "Николай Лесков": ["Николай Лесков", "Лесков"],
    "Максим Горький": ["Максим Горький", "Горький"],
    "Иван Бунин": ["Иван Бунин", "Бунин"],
    "Александр Куприн": ["Александр Куприн", "Куприн"],
}

# Works by Maugham are not yet public domain in Finland/EU (death 1965).
EXCLUDED_AUTHORS = {"Сомерсет Моэм"}
LOCAL_AUTHORS = {"Иван Тургенев", "Николай Лесков", "Максим Горький", "Иван Бунин", "Александр Куприн"}
LOCAL_REPO = Path(os.environ.get("DANKERTLIT_DIR", "/tmp/dankertlit"))


def clean(s: str) -> str:
    s = html.unescape(s)
    s = s.replace("\xa0", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def api(params):
    params = dict(params)
    params["format"] = "json"
    params["formatversion"] = 2
    r = requests.get(API, params=params, headers={"User-Agent": UA}, timeout=45)
    r.raise_for_status()
    return r.json()


def candidate_titles(story):
    title = story["title"]
    author = story["author"]
    queries = [
        f'intitle:"{title}" "{AUTHOR_ALIASES[author][0]}"',
        f'"{title}" "{AUTHOR_ALIASES[author][-1]}"',
        f'intitle:"{title}"',
    ]
    found = []
    for q in queries:
        data = api({
            "action": "query",
            "list": "search",
            "srsearch": q,
            "srnamespace": 0,
            "srlimit": 10,
        })
        for row in data.get("query", {}).get("search", []):
            t = row["title"]
            if t not in found:
                found.append(t)
        if found:
            break
        time.sleep(0.15)
    return found


def score_candidate(story, page_title, text):
    title_words = set(re.findall(r"[а-яёa-z0-9]+", story["title"].lower()))
    page_words = set(re.findall(r"[а-яёa-z0-9]+", page_title.lower()))
    overlap = len(title_words & page_words)
    author_hit = any(a.lower() in (page_title + " " + text[:2500]).lower()
                     for a in AUTHOR_ALIASES[story["author"]])
    bad = any(x in page_title.lower() for x in ["автор:", "категория:", "обсуждение:", "указатель"])
    return overlap * 3 + (5 if author_hit else 0) - (20 if bad else 0)


def extract_page(page_title):
    data = api({
        "action": "parse",
        "page": page_title,
        "prop": "text|displaytitle|categories",
        "disableeditsection": 1,
        "redirects": 1,
    })
    parsed = data.get("parse")
    if not parsed:
        return None
    soup = BeautifulSoup(parsed["text"], "html.parser")
    root = soup.select_one(".mw-parser-output") or soup
    for sel in [
        "table", ".navbox", ".metadata", ".ambox", ".infobox", ".toc",
        ".mw-editsection", ".noprint", ".printfooter", ".catlinks",
        "sup.reference", "ol.references", ".licenseContainer",
        ".ws-noexport", "style", "script",
    ]:
        for node in root.select(sel):
            node.decompose()
    # Convert headings and paragraphs into a stable plain-text representation.
    parts = []
    for node in root.find_all(["h2", "h3", "h4", "p", "div", "blockquote"], recursive=True):
        if node.find_parent(["table", "blockquote"]) and node.name != "blockquote":
            continue
        txt = clean(node.get_text(" ", strip=True))
        if not txt or len(txt) < 2:
            continue
        if txt in parts[-3:]:
            continue
        if node.name in ("h2", "h3", "h4"):
            parts.append("\n" + txt + "\n")
        elif len(txt) >= 25:
            parts.append(txt)
    text = clean("\n\n".join(parts))
    url = "https://ru.wikisource.org/wiki/" + quote(page_title.replace(" ", "_"), safe="()/,:;")
    cats = [c.get("*", "") for c in parsed.get("categories", [])]
    return {"title": page_title, "text": text, "url": url, "categories": cats}


def norm_name(value: str) -> str:
    value = value.lower().replace("ё", "е").replace("…", "...")
    value = re.sub(r"[^\wа-яa-z0-9]+", " ", value, flags=re.I)
    return re.sub(r"\s+", " ", value).strip()


def local_story_dir(story):
    author_dir = LOCAL_REPO / "lit" / story["author"]
    if not author_dir.exists():
        return None
    target = norm_name(story["title"])
    exact = author_dir / story["title"]
    if exact.exists() and exact.is_dir():
        return exact
    best = None
    best_score = 0
    target_words = set(target.split())
    for child in author_dir.iterdir():
        if not child.is_dir():
            continue
        candidate = norm_name(child.name)
        if candidate == target:
            return child
        words = set(candidate.split())
        overlap = len(words & target_words)
        score = overlap / max(len(target_words), 1)
        if score > best_score:
            best_score, best = score, child
    return best if best_score >= 0.8 else None


def local_extract(story):
    folder = local_story_dir(story)
    if not folder:
        return None
    files = []
    root_main = folder / "main.md"
    if root_main.exists():
        files.append(root_main)
    def part_num(path):
        m = re.search(r"(\d+)", path.parent.name)
        return int(m.group(1)) if m else 10**9
    files.extend(sorted(folder.glob("часть */main.md"), key=part_num))
    if not files:
        files.extend(sorted(folder.glob("*/main.md"), key=part_num))
    blocks = []
    for f in files:
        txt = clean(f.read_text("utf-8", errors="ignore"))
        if txt:
            blocks.append(txt)
    text = clean("\n\n".join(blocks))
    if len(text) < 800:
        return None
    source_url = ""
    data_file = folder / "data.json"
    if data_file.exists():
        try:
            meta = json.loads(data_file.read_text("utf-8"))
            source_url = meta.get("url", "")
        except Exception:
            pass
    if not source_url:
        source_url = "https://github.com/DanKertLit/dankertlit.github.io"
    return {"title": story["title"], "text": text, "url": source_url}


def find_story(story):
    best = None
    for page_title in candidate_titles(story):
        try:
            page = extract_page(page_title)
        except Exception:
            continue
        if not page or len(page["text"]) < 1200:
            continue
        sc = score_candidate(story, page_title, page["text"])
        if best is None or sc > best[0]:
            best = (sc, page)
        time.sleep(0.1)
    if best and best[0] >= 3:
        return best[1]
    return None


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:" + edge
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, value in kwargs[edge].items():
                element.set(qn("w:" + key), str(value))


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_doc(found, missing, output):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.1)
    sec.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Liberation Serif"
    styles["Normal"].font.size = Pt(11.5)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    for st in ["Title", "Heading 1", "Heading 2"]:
        styles[st].font.name = "Liberation Serif"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("МАЛОИЗВЕСТНЫЕ РАССКАЗЫ ИМЕНИТЫХ ПИСАТЕЛЕЙ")
    r.bold = True
    r.font.size = Pt(19)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Полные тексты, доступные на русской Викитеке")
    r.italic = True
    r.font.size = Pt(13)

    doc.add_paragraph(
        ("В сборник включены только тексты с подтверждённо свободных источников. "
         "Для русских авторов использованы полные оригинальные тексты из цифровой коллекции, "
         "собранной с iLibrary; для переводных произведений — только найденные на русской Викитеке "
         "свободные версии. Неподтверждённые переводы пропущены. Амброз Бирс исключён по условию подборки.")
    )
    doc.add_paragraph(
        f"Найдено и включено: {len(found)}. Не включено: {len(missing)}. "
        "У каждого рассказа указана страница-источник."
    )

    grouped = OrderedDict()
    for item in found:
        grouped.setdefault(item["story"]["author"], []).append(item)

    doc.add_heading("Содержание сборника", level=1)
    for author, items in grouped.items():
        doc.add_paragraph(f"{author} — {len(items)} текст(ов)", style="List Bullet")

    for author, items in grouped.items():
        doc.add_page_break()
        doc.add_heading(author, level=1)
        for item in items:
            story, page = item["story"], item["page"]
            title = story["title"]
            if story.get("original"):
                title += f" ({story['original']})"
            doc.add_heading(title, level=2)
            src = doc.add_paragraph()
            src.add_run("Источник: ").bold = True
            add_hyperlink(src, page["title"] + " — Викитека", page["url"])
            for block in re.split(r"\n\s*\n", page["text"]):
                block = clean(block)
                if not block:
                    continue
                p = doc.add_paragraph(block)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(0.7)

    doc.add_page_break()
    doc.add_heading("Не включённые произведения", level=1)
    doc.add_paragraph(
        "Эти позиции отсутствовали на русской Викитеке под проверенными названиями, "
        "оказались слишком короткими служебными страницами либо не были включены по правовому статусу автора."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Автор"
    hdr[1].text = "Произведение"
    hdr[2].text = "Причина"
    for row in missing:
        cells = table.add_row().cells
        cells[0].text = row["author"]
        cells[1].text = row["title"]
        cells[2].text = row["reason"]
    doc.save(output)


def main():
    outdir = Path("output")
    outdir.mkdir(exist_ok=True)
    found = []
    missing = []
    cache_path = outdir / "results.json"
    if cache_path.exists():
        state = json.loads(cache_path.read_text("utf-8"))
        found = state.get("found", [])
        missing = state.get("missing", [])
    else:
        for i, story in enumerate(STORIES, 1):
            print(f"[{i}/{len(STORIES)}] {story['author']} — {story['title']}", flush=True)
            if story["author"] in EXCLUDED_AUTHORS:
                missing.append({**story, "reason": "автор ещё не перешёл в общественное достояние в Финляндии/ЕС"})
                continue
            try:
                if story["author"] in LOCAL_AUTHORS:
                    page = local_extract(story)
                    if not page:
                        page = find_story(story)
                else:
                    page = find_story(story)
            except Exception as exc:
                missing.append({**story, "reason": f"ошибка поиска: {exc}"})
                continue
            if page:
                found.append({"story": story, "page": page})
            else:
                missing.append({**story, "reason": "не найден подтверждённый полный свободный текст на русской Викитеке"})
        cache_path.write_text(json.dumps({"found": found, "missing": missing}, ensure_ascii=False, indent=2), "utf-8")

    build_doc(found, missing, outdir / "rasskazy_obshchestvennoe_dostoyanie.docx")
    (outdir / "README.txt").write_text(
        f"Включено полных текстов: {len(found)}\nПропущено: {len(missing)}\n"
        "Источники включённых текстов: iLibrary/цифровая коллекция и русская Викитека.\n",
        "utf-8"
    )
    print(f"DONE: found={len(found)}, missing={len(missing)}", flush=True)


if __name__ == "__main__":
    main()
